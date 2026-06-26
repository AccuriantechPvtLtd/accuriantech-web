"""
Build a clean, ATS-friendly, client-submission resume .docx for PARAG
- Removes last name (first name only — document only shows "PARAG")
- Removes company names from Professional Experience and Project sections
- Adds Accuriantech logo on its own line (top-right), name on separate line below
- Preserves all content, durations, roles, responsibilities, technical stack
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LOGO_PATH = os.path.join(os.path.dirname(__file__), "public", "logo.png")
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "Parag_Resume.docx")

NAVY = RGBColor(0x0B, 0x1F, 0x4D)
SLATE = RGBColor(0x33, 0x33, 0x33)
MUTED = RGBColor(0x55, 0x5F, 0x6D)

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = SLATE


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            border = OxmlElement(f"w:{edge}")
            for k, v in kwargs[edge].items():
                border.set(qn(f"w:{k}"), str(v))
            tc_borders.append(border)
    tc_pr.append(tc_borders)


def add_section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0B1F4D")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return p


def add_role_block(title, duration):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.6), WD_ALIGN_PARAGRAPH.RIGHT)
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = NAVY
    p.add_run("\t")
    run_dur = p.add_run(duration)
    run_dur.italic = True
    run_dur.font.size = Pt(10)
    run_dur.font.color.rgb = MUTED


def add_subline(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        run.font.color.rgb = SLATE


def add_skills_table(rows_data):
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.6)
    for i, (category, skills) in enumerate(rows_data):
        cat_cell = table.cell(i, 0)
        skill_cell = table.cell(i, 1)
        cat_cell.width = Inches(2.0)
        skill_cell.width = Inches(4.6)
        cat_p = cat_cell.paragraphs[0]
        cat_p.paragraph_format.space_after = Pt(2)
        cat_run = cat_p.add_run(category)
        cat_run.bold = True
        cat_run.font.size = Pt(10)
        cat_run.font.color.rgb = NAVY
        sk_p = skill_cell.paragraphs[0]
        sk_p.paragraph_format.space_after = Pt(2)
        sk_run = sk_p.add_run(skills)
        sk_run.font.size = Pt(10)
        sk_run.font.color.rgb = SLATE
        for cell in (cat_cell, skill_cell):
            set_cell_border(cell, top={"val": "nil"}, left={"val": "nil"},
                          bottom={"val": "single", "sz": "4", "color": "D9D9D9"}, right={"val": "nil"})


# ============================================================
# HEADER
# ============================================================
logo_para = doc.add_paragraph()
logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
logo_para.paragraph_format.space_before = Pt(0)
logo_para.paragraph_format.space_after = Pt(6)
logo_run = logo_para.add_run()
if os.path.exists(LOGO_PATH):
    logo_run.add_picture(LOGO_PATH, width=Inches(1.4))

name_para = doc.add_paragraph()
name_para.paragraph_format.space_before = Pt(0)
name_para.paragraph_format.space_after = Pt(2)
name_run = name_para.add_run("PARAG")
name_run.bold = True
name_run.font.size = Pt(26)
name_run.font.color.rgb = NAVY
name_run.font.name = "Calibri"

tag_para = doc.add_paragraph()
tag_para.paragraph_format.space_before = Pt(0)
tag_para.paragraph_format.space_after = Pt(0)
tag_run = tag_para.add_run(
    "Solution Architect  |  Data Cloud & AI Consultant  |  Customer 360 Architect  |  Technical Lead"
)
tag_run.font.size = Pt(10.5)
tag_run.font.color.rgb = MUTED
tag_run.italic = True

tag2_para = doc.add_paragraph()
tag2_para.paragraph_format.space_before = Pt(0)
tag2_para.paragraph_format.space_after = Pt(0)
tag2_run = tag2_para.add_run("10+ Years of Experience  |  14X Salesforce Certified")
tag2_run.font.size = Pt(10.5)
tag2_run.font.color.rgb = MUTED
tag2_run.italic = True

# Divider
div = doc.add_paragraph()
div.paragraph_format.space_before = Pt(2)
div.paragraph_format.space_after = Pt(0)
p_pr = div._p.get_or_add_pPr()
p_bdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "12")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), "0B1F4D")
p_bdr.append(bottom)
p_pr.append(p_bdr)

# ============================================================
# PROFESSIONAL SUMMARY
# ============================================================
add_section_heading("Professional Summary")
add_bullets([
    "Salesforce Solution Architect and Technical Lead with 10+ years of experience delivering enterprise Salesforce solutions across Sales Cloud, Service Cloud, Experience Cloud, Marketing Cloud, and Salesforce Data Cloud for large-scale digital transformation initiatives.",
    "Strong expertise in architecting Customer 360 ecosystems by integrating CRM, marketing, and enterprise platforms to deliver unified customer experiences, AI-driven business solutions, and scalable customer engagement strategies.",
    "Hands-on experience designing and implementing Salesforce AI and Agentforce solutions, including autonomous Sales and Service Agents, intelligent workflow automation, prompt engineering, and AI-assisted customer interactions.",
    "Extensive experience designing API-led integration architectures using MuleSoft, REST/SOAP APIs, Apex Callouts, Platform Events, asynchronous processing patterns, and enterprise integration frameworks across complex multi-system environments.",
    "Strong knowledge of Salesforce Data Cloud, Customer Data Platform (CDP), Data Streams, Identity Resolution, Unified Customer Profiles, Calculated Insights, Audience Segmentation, Customer Activation, and Data 360 concepts for enterprise customer data management.",
    "Hands-on expertise in Apex, Lightning Web Components (LWC), Salesforce Flow, OmniStudio, Salesforce DX, CI/CD automation, enterprise security architecture, and scalable Salesforce application development.",
    "Experience working across Media and Marketing domains, supporting customer acquisition, campaign management, audience activation, personalization, and AI-enabled engagement initiatives while collaborating with enterprise architecture teams to define scalable technical solutions.",
    "Proven leader with strong experience conducting architecture and code reviews, mentoring technical teams, establishing development standards, and driving Agile delivery practices while continuously upskilling in emerging Salesforce AI and automation technologies.",
])

# ============================================================
# CERTIFICATIONS
# ============================================================
add_section_heading("Certifications")
add_bullets([
    "Salesforce Certified Administrator",
    "Salesforce Platform Developer I",
    "Salesforce Platform Developer II",
    "Salesforce Certified Sales Cloud Consultant",
    "Salesforce Certified Service Cloud Consultant",
    "Salesforce Certified Experience Cloud Consultant",
    "Salesforce Certified Platform App Builder",
    "Salesforce Certified Data Architect",
    "Salesforce Certified Sharing & Visibility Architect",
    "Salesforce Certified Application Architect",
    "Salesforce Certified AI Associate",
    "Salesforce Certified AI Specialist",
    "Einstein Analytics and Discovery Consultant",
    "Copado Fundamentals I",
    "Copado Fundamentals II",
])

# ============================================================
# TECHNICAL SKILLS
# ============================================================
add_section_heading("Technical Skills")
add_skills_table([
    ("Architecture & CRM", "Salesforce Solution Architecture, Application Architecture, Enterprise CRM Transformation, Customer 360 Architecture, Multi-Org Strategy, Salesforce Governance Framework, Security & Sharing Architecture, Enterprise Data Modeling, Agile & SCRUM Delivery"),
    ("Salesforce Clouds", "Sales Cloud, Service Cloud, Experience Cloud, Marketing Cloud, Salesforce Data Cloud, Customer Data Platform (CDP), Salesforce Einstein Analytics, OmniStudio, Salesforce CPQ"),
    ("AI & Agentforce", "Salesforce Agentforce, AI-Powered Sales & Service Automation, Intelligent Agent Design, Prompt Engineering, Einstein AI Capabilities, AI-Driven Customer Engagement, Agent Actions & Topics Configuration, Generative AI Use Cases, AI Governance & Responsible AI, RAG, Grounded Prompt Templates, Agentic AI Workflows"),
    ("Data Cloud & Customer 360", "Data Streams, Data Model Objects (DMOs), Identity Resolution, Unified Customer Profiles, Golden Customer Record, Calculated Insights, Audience Segmentation, Customer Activation, Consent Management, Data Harmonization, Customer Profile Stitching, Data Governance, Metadata Management"),
    ("Integration & APIs", "REST APIs, SOAP APIs, MuleSoft, Apex Callouts, Platform Events, Event-Driven Architecture, Asynchronous Processing, External Services, Pub/Sub Integration Patterns, Enterprise Integration Frameworks, Third-Party System Integration"),
    ("Salesforce Development", "Apex Classes, Apex Triggers, Lightning Web Components (LWC), Aura Components, Visualforce, SOQL, SOSL, Salesforce Flow, Process Automation, Validation Rules, Approval Processes"),
    ("DevOps & Release", "Salesforce DX, Copado, Git, GitHub, Jenkins, Azure DevOps, CI/CD Pipeline Design, Release Governance, Deployment Automation, Environment Management"),
    ("Data & Analytics", "SQL, Snowflake Integration, Data Migration, Data Transformation, Data Quality Management, Data Loader, Reports & Dashboards, Tableau CRM, Customer Analytics, Marketing Analytics"),
    ("Domain Expertise", "Media & Entertainment, Digital Marketing, Customer Engagement, Campaign Management, Lead-to-Revenue Processes, Customer Support Operations, Enterprise CRM Modernization, Personalization & Audience Activation"),
])


# ============================================================
# PROFESSIONAL EXPERIENCE (company names removed)
# ============================================================
add_section_heading("Professional Experience")

add_role_block("Technical Lead / Salesforce Solution Architect", "March 2023 – Present")
add_subline("Responsibilities:")
add_bullets([
    "Lead enterprise Salesforce architecture initiatives focused on AI-driven customer engagement, Customer 360 transformation, and intelligent automation solutions.",
    "Design scalable Salesforce ecosystems integrating Sales Cloud, Service Cloud, Marketing Cloud, Experience Cloud, and Data Cloud.",
    "Architect Agentforce-based business solutions to automate customer interactions, service operations, and sales assistance workflows.",
    "Build intelligent agent capabilities using Salesforce AI, prompt engineering strategies, and enterprise knowledge sources.",
    "Design unified customer data models using Data Cloud, identity resolution, and Golden Customer Profile architecture.",
    "Define enterprise integration architecture standards leveraging MuleSoft, API-led connectivity, Apex Callouts, Platform Events, and asynchronous messaging frameworks.",
    "Integrate Salesforce with marketing platforms, analytics systems, and external enterprise applications to enable seamless customer experiences.",
    "Collaborate with business stakeholders to define AI use cases, customer journey automation strategies, and enterprise architecture roadmaps.",
    "Conduct architecture reviews, code quality assessments, and technical governance sessions.",
    "Mentor development teams on modern Salesforce development patterns, AI capabilities, and enterprise integration standards.",
    "Drive Agile delivery practices and support end-to-end solution implementation from design through production deployment.",
])

add_role_block("Associate Technical Lead", "Jan 2022 – Feb 2023")
add_subline("Responsibilities:")
add_bullets([
    "Led technical design activities for enterprise Salesforce transformation programs involving Sales Cloud, Service Cloud, and Data Cloud integrations.",
    "Developed scalable integration frameworks connecting Salesforce with external business applications through REST APIs and MuleSoft services.",
    "Supported enterprise customer data unification initiatives using Data Cloud ingestion and identity resolution capabilities.",
    "Designed reusable Apex services, Lightning Web Components, and Flow-based automation frameworks.",
    "Implemented CI/CD pipelines using Copado, Salesforce DX, Git, and automated deployment strategies.",
    "Worked closely with solution architects and business teams to deliver scalable CRM solutions for marketing and customer engagement operations.",
    "Guided junior developers and conducted technical design and code review sessions.",
])

add_role_block("Senior Salesforce Developer", "Oct 2020 – Jan 2022")
add_subline("Responsibilities:")
add_bullets([
    "Delivered enterprise Salesforce solutions across Sales Cloud and Service Cloud for global business customers.",
    "Built Apex services, Lightning Web Components, Flows, and reusable integration components.",
    "Implemented secure API integrations between Salesforce and third-party enterprise systems.",
    "Participated in enterprise architecture discussions around customer data management and CRM modernization.",
    "Enhanced application performance through optimized SOQL strategies and asynchronous processing techniques.",
    "Designed extensible customer data models supporting future analytics and AI initiatives.",
    "Collaborated with architects and business teams throughout Agile delivery cycles.",
])

add_role_block("Software Engineer", "Jan 2020 – Oct 2020")
add_subline("Responsibilities:")
add_bullets([
    "Developed Salesforce applications using Apex, LWC, Visualforce, and declarative automation.",
    "Built service management workflows and customer support automation solutions.",
    "Integrated Salesforce with external platforms using REST APIs and middleware services.",
    "Participated in enterprise data migration and customer data quality initiatives.",
    "Supported deployment activities and release management processes.",
])

add_role_block("Associate Cloud Consultant", "Dec 2018 – Dec 2019")
add_subline("Responsibilities:")
add_bullets([
    "Delivered CRM modernization initiatives across Sales Cloud and Service Cloud environments.",
    "Configured automation processes, security models, reports, and dashboards.",
    "Assisted with customer data migration and integration activities.",
    "Supported business process optimization and stakeholder workshops.",
    "Participated in Agile project execution and deployment planning.",
])

add_role_block("Salesforce Junior Developer", "Sept 2016 – Nov 2018")
add_subline("Responsibilities:")
add_bullets([
    "Developed custom Salesforce applications using Apex, Visualforce, and Aura Components.",
    "Built workflow automation and business process solutions.",
    "Supported REST API integrations and external system connectivity.",
    "Assisted with deployment coordination and production support activities.",
    "Participated in data migration and metadata management tasks.",
])

add_role_block("Salesforce Intern", "March 2016 – Aug 2016")
add_subline("Responsibilities:")
add_bullets([
    "Assisted in Salesforce application development and configuration activities.",
    "Worked on Sales Cloud customization and automation requirements.",
    "Supported data migration, validation, and testing efforts.",
    "Participated in integration support and deployment validation activities.",
    "Collaborated with senior developers during the complete software development lifecycle.",
])

# ============================================================
# PROJECT EXPERIENCE
# ============================================================
add_section_heading("Project Experience")


def add_project(title, role, duration, skills, responsibilities):
    add_role_block(title, duration)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("Role: ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(role)
    r2.font.size = Pt(10)
    r2.font.color.rgb = SLATE
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r3 = p2.add_run("Skills Used: ")
    r3.bold = True
    r3.font.size = Pt(10)
    r3.font.color.rgb = NAVY
    r4 = p2.add_run(skills)
    r4.font.size = Pt(10)
    r4.font.color.rgb = SLATE
    add_subline("Responsibilities:")
    add_bullets(responsibilities)


add_project(
    "Enterprise Agentforce Sales & Service Automation Platform",
    "Salesforce Solution Architect",
    "Jan 2025 – Present",
    "Salesforce Agentforce, Salesforce AI, Data Cloud, Sales Cloud, Service Cloud, Prompt Engineering, Apex, LWC, MuleSoft, REST APIs, Platform Events",
    [
        "Architected an AI-powered enterprise solution leveraging Salesforce Agentforce to automate sales assistance and customer service operations.",
        "Designed and implemented autonomous Sales and Service Agents leveraging Agentforce, capable of executing contextual business actions using CRM and enterprise knowledge sources.",
        "Configured trusted topics, custom agent actions, and AI prompt templates aligned with business processes.",
        "Integrated Agentforce with Salesforce CRM and external enterprise systems through MuleSoft and REST APIs.",
        "Utilized Data Cloud unified customer profiles to provide contextual intelligence for AI interactions.",
        "Built reusable Apex services and LWC components supporting AI-assisted workflows.",
        "Worked closely with business stakeholders to identify automation opportunities and optimize customer journeys.",
        "Led architecture reviews and mentored technical teams on Salesforce AI capabilities.",
    ],
)

add_project(
    "Enterprise Customer 360 & Data Cloud Transformation Platform",
    "Salesforce Solution Architect",
    "Mar 2023 – Dec 2024",
    "Salesforce Data Cloud, Customer 360, Identity Resolution, Data Streams, DMOs, Calculated Insights, Audience Segmentation, Snowflake, REST APIs",
    [
        "Designed enterprise Customer 360 architecture for consolidating customer data across multiple business platforms.",
        "Supported federated customer data access patterns to enable AI-driven activation across enterprise systems.",
        "Implemented Data Streams and Data Model Objects (DMOs) for batch and real-time customer data ingestion.",
        "Configured identity resolution rules to create unified customer profiles.",
        "Built calculated insights and audience segmentation models for personalized engagement.",
        "Integrated Data Cloud with Sales Cloud, Service Cloud, Marketing Cloud, and enterprise data warehouses.",
        "Established enterprise data governance and metadata management standards.",
        "Collaborated with marketing and analytics teams to improve customer activation strategies.",
        "Supported AI-driven use cases through centralized customer intelligence.",
    ],
)

add_project(
    "Media & Marketing Customer Engagement Platform",
    "Salesforce Technical Architect",
    "Jul 2022 – Feb 2023",
    "Marketing Cloud, Sales Cloud, Data Cloud, Journey Builder, Contact Builder, REST APIs, Audience Segmentation",
    [
        "Designed Salesforce architecture supporting digital marketing and media campaign operations.",
        "Integrated Marketing Cloud with Sales Cloud to provide a unified customer engagement experience.",
        "Built audience segmentation models using behavioral and CRM data.",
        "Developed API integrations between marketing platforms and external content management systems.",
        "Supported personalized customer journeys using customer profile data.",
        "Improved campaign visibility through custom dashboards and analytics.",
        "Worked with business teams to optimize lead-to-revenue processes.",
        "Assisted in technical governance and deployment planning activities.",
    ],
)

add_project(
    "Enterprise CRM Integration & API Modernization",
    "Associate Technical Lead",
    "Jan 2022 – Jun 2022",
    "MuleSoft, REST APIs, SOAP APIs, Apex Callouts, Platform Events, Salesforce DX, CI/CD",
    [
        "Designed API-led integration architecture connecting Salesforce with multiple enterprise applications.",
        "Developed secure REST and SOAP integrations using MuleSoft and Apex Callouts.",
        "Implemented asynchronous processing and Platform Event-based communication models.",
        "Created reusable integration services to support multiple business functions.",
        "Improved system reliability through standardized error handling and monitoring frameworks.",
        "Worked with integration teams to simplify cross-platform data exchange.",
        "Supported deployment automation using Copado and Salesforce DX.",
        "Participated in architecture reviews and technical design sessions.",
    ],
)

add_project(
    "Global CRM Platform Modernization Program",
    "Senior Salesforce Developer",
    "Oct 2020 – Jan 2022",
    "Sales Cloud, Service Cloud, Apex, LWC, Salesforce Flow, REST APIs, CI/CD",
    [
        "Delivered scalable Salesforce solutions supporting global sales and customer support operations.",
        "Developed custom business logic using Apex classes and triggers.",
        "Built Lightning Web Components for enhanced user experience.",
        "Implemented workflow automation using Salesforce Flows.",
        "Integrated Salesforce with enterprise applications through REST APIs.",
        "Supported customer data migration and platform optimization initiatives.",
        "Participated in Agile ceremonies and release planning activities.",
        "Collaborated with architects to improve application scalability.",
    ],
)

add_project(
    "Insurance Claims & Customer Service Automation",
    "Software Engineer",
    "Jan 2020 – Oct 2020",
    "Service Cloud, Apex, Lightning Components, Salesforce Flow, REST APIs, Data Migration",
    [
        "Developed Service Cloud applications supporting insurance claims processing.",
        "Built Apex classes and Lightning Components for case management automation.",
        "Configured escalation processes and service workflows.",
        "Integrated Salesforce with external policy administration systems.",
        "Assisted in customer data migration and validation activities.",
        "Supported deployment coordination and production issue resolution.",
        "Created operational reports and dashboards.",
        "Participated in testing and user acceptance activities.",
    ],
)

add_project(
    "Enterprise CRM Modernization & Sales Automation",
    "Associate Cloud Consultant",
    "Dec 2018 – Dec 2019",
    "Sales Cloud, Service Cloud, Apex, Reports & Dashboards, Security & Sharing, Data Loader",
    [
        "Implemented CRM solutions supporting enterprise sales and customer management operations.",
        "Developed automation using Apex and declarative Salesforce features.",
        "Configured security models, profiles, permission sets, and sharing rules.",
        "Supported customer data migration and cleansing activities.",
        "Created operational reports and executive dashboards.",
        "Assisted with deployment planning and metadata management.",
        "Worked with business users during requirement gathering sessions.",
        "Supported post-production enhancements and maintenance.",
    ],
)

add_project(
    "Customer Service & Sales Management Solution",
    "Salesforce Developer",
    "Sept 2016 – Nov 2018",
    "Sales Cloud, Service Cloud, Apex, Visualforce, Aura Components, SOQL, REST APIs",
    [
        "Developed custom Salesforce applications for customer and sales management processes.",
        "Created Apex classes, triggers, and Visualforce pages to automate business operations.",
        "Built Aura Components to improve application usability.",
        "Integrated Salesforce with external applications using REST APIs.",
        "Developed validation rules and approval processes.",
        "Assisted with bulk data migration and Data Loader activities.",
        "Supported testing, deployment, and production support functions.",
        "Collaborated with functional teams to deliver business requirements.",
    ],
)

# Save
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
